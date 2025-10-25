# File: pqc_inspector_server/services/document_processor.py
# 📄 PDF, DOCX, TXT 등 다양한 문서 형식을 처리하는 서비스입니다.

import os
import re
from typing import List, Dict, Any, Optional
from pathlib import Path
import hashlib

# PDF 처리용
try:
    import fitz  # PyMuPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# DOCX 처리용
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# 마크다운 처리용
try:
    import markdown
    from bs4 import BeautifulSoup
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False

class DocumentProcessor:
    def __init__(self):
        self.supported_formats = []
        self.chunk_size = 1000  # 기본 청크 크기 (문자 수)
        self.chunk_overlap = 100  # 청크 간 중복 문자 수

        # 지원 가능한 형식 확인
        if PDF_AVAILABLE:
            self.supported_formats.extend(['.pdf'])
        if DOCX_AVAILABLE:
            self.supported_formats.extend(['.docx'])
        if MARKDOWN_AVAILABLE:
            self.supported_formats.extend(['.md', '.markdown'])

        self.supported_formats.extend(['.txt', '.py', '.c', '.cpp', '.h', '.js', '.java'])

        print(f"DocumentProcessor 초기화됨. 지원 형식: {self.supported_formats}")

    def is_supported(self, file_path: str) -> bool:
        """파일 형식이 지원되는지 확인합니다."""
        suffix = Path(file_path).suffix.lower()
        return suffix in self.supported_formats

    async def process_document(self, file_path: str, agent_type: str = "auto") -> Dict[str, Any]:
        """
        문서를 처리하여 RAG 시스템에 추가 가능한 형태로 변환합니다.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"파일을 찾을 수 없습니다: {file_path}")

        if not self.is_supported(file_path):
            raise ValueError(f"지원하지 않는 파일 형식: {Path(file_path).suffix}")

        print(f"📄 문서 처리 시작: {file_path}")

        # 파일 정보 추출
        file_info = self._extract_file_info(file_path)

        # 텍스트 추출
        raw_text = self._extract_text(file_path)

        # 전처리
        processed_text = self._preprocess_text(raw_text)

        # 청킹
        chunks = self._chunk_text(processed_text)

        # 에이전트 타입 자동 결정
        if agent_type == "auto":
            agent_type = self._determine_agent_type(file_path, processed_text)

        # 메타데이터 생성
        result = {
            "file_info": file_info,
            "agent_type": agent_type,
            "chunks": chunks,
            "total_chunks": len(chunks),
            "original_text_length": len(raw_text),
            "processed_text_length": len(processed_text)
        }

        print(f"✅ 문서 처리 완료: {len(chunks)}개 청크 생성")
        return result

    def _extract_file_info(self, file_path: str) -> Dict[str, Any]:
        """파일 기본 정보를 추출합니다."""
        path_obj = Path(file_path)
        stat = path_obj.stat()

        with open(file_path, 'rb') as f:
            file_hash = hashlib.md5(f.read()).hexdigest()

        return {
            "filename": path_obj.name,
            "file_size": stat.st_size,
            "file_extension": path_obj.suffix.lower(),
            "file_hash": file_hash,
            "creation_time": stat.st_ctime,
            "modification_time": stat.st_mtime
        }

    def _extract_text(self, file_path: str) -> str:
        """파일 형식에 따라 텍스트를 추출합니다."""
        suffix = Path(file_path).suffix.lower()

        if suffix == '.pdf':
            return self._extract_pdf_text(file_path)
        elif suffix == '.docx':
            return self._extract_docx_text(file_path)
        elif suffix in ['.md', '.markdown']:
            return self._extract_markdown_text(file_path)
        else:
            # 일반 텍스트 파일
            return self._extract_plain_text(file_path)

    def _extract_pdf_text(self, file_path: str) -> str:
        """PDF에서 텍스트를 추출합니다."""
        if not PDF_AVAILABLE:
            raise ImportError("PDF 처리를 위해 PyMuPDF를 설치하세요: pip install PyMuPDF")

        text = ""
        try:
            doc = fitz.open(file_path)
            for page_num in range(len(doc)):
                page = doc[page_num]
                page_text = page.get_text()

                # 페이지 번호 추가
                text += f"\n--- 페이지 {page_num + 1} ---\n"
                text += page_text

            doc.close()
            print(f"PDF 텍스트 추출 완료: {len(doc)} 페이지")
            return text

        except Exception as e:
            raise Exception(f"PDF 텍스트 추출 실패: {e}")

    def _extract_docx_text(self, file_path: str) -> str:
        """DOCX에서 텍스트를 추출합니다."""
        if not DOCX_AVAILABLE:
            raise ImportError("DOCX 처리를 위해 python-docx를 설치하세요: pip install python-docx")

        try:
            doc = Document(file_path)
            text = ""

            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            print(f"DOCX 텍스트 추출 완료: {len(doc.paragraphs)} 문단")
            return text

        except Exception as e:
            raise Exception(f"DOCX 텍스트 추출 실패: {e}")

    def _extract_markdown_text(self, file_path: str) -> str:
        """마크다운에서 텍스트를 추출합니다."""
        if not MARKDOWN_AVAILABLE:
            raise ImportError("마크다운 처리를 위해 markdown과 beautifulsoup4를 설치하세요")

        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                md_content = f.read()

            # 마크다운을 HTML로 변환 후 텍스트 추출
            html = markdown.markdown(md_content)
            soup = BeautifulSoup(html, 'html.parser')
            text = soup.get_text()

            print(f"Markdown 텍스트 추출 완료")
            return text

        except Exception as e:
            raise Exception(f"Markdown 텍스트 추출 실패: {e}")

    def _extract_plain_text(self, file_path: str) -> str:
        """일반 텍스트 파일에서 텍스트를 추출합니다."""
        try:
            encodings = ['utf-8', 'utf-16', 'cp949', 'euc-kr', 'latin-1']

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    print(f"텍스트 파일 읽기 완료 (인코딩: {encoding})")
                    return text
                except UnicodeDecodeError:
                    continue

            raise Exception("지원하는 인코딩으로 파일을 읽을 수 없습니다")

        except Exception as e:
            raise Exception(f"텍스트 파일 읽기 실패: {e}")

    def _preprocess_text(self, text: str) -> str:
        """텍스트를 전처리합니다."""
        # 과도한 공백 제거
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r' +', ' ', text)

        # 특수 문자 정리 (필요시)
        text = text.strip()

        return text

    def _chunk_text(self, text: str) -> List[Dict[str, Any]]:
        """텍스트를 청크로 분할합니다."""
        chunks = []

        # 문단 기준으로 1차 분할
        paragraphs = text.split('\n\n')

        current_chunk = ""
        chunk_index = 0

        for paragraph in paragraphs:
            # 현재 청크에 문단을 추가했을 때 크기 확인
            test_chunk = current_chunk + "\n\n" + paragraph if current_chunk else paragraph

            if len(test_chunk) <= self.chunk_size:
                current_chunk = test_chunk
            else:
                # 현재 청크가 비어있지 않으면 저장
                if current_chunk:
                    chunks.append(self._create_chunk(current_chunk, chunk_index))
                    chunk_index += 1

                # 단일 문단이 너무 클 경우 강제 분할
                if len(paragraph) > self.chunk_size:
                    sub_chunks = self._force_split_text(paragraph, chunk_index)
                    chunks.extend(sub_chunks)
                    chunk_index += len(sub_chunks)
                    current_chunk = ""
                else:
                    current_chunk = paragraph

        # 마지막 청크 처리
        if current_chunk:
            chunks.append(self._create_chunk(current_chunk, chunk_index))

        return chunks

    def _create_chunk(self, text: str, index: int) -> Dict[str, Any]:
        """청크 객체를 생성합니다."""
        return {
            "index": index,
            "content": text.strip(),
            "char_count": len(text),
            "word_count": len(text.split()),
            "chunk_hash": hashlib.md5(text.encode()).hexdigest()[:8]
        }

    def _force_split_text(self, text: str, start_index: int) -> List[Dict[str, Any]]:
        """긴 텍스트를 강제로 분할합니다."""
        chunks = []
        index = start_index

        for i in range(0, len(text), self.chunk_size - self.chunk_overlap):
            chunk_text = text[i:i + self.chunk_size]
            chunks.append(self._create_chunk(chunk_text, index))
            index += 1

        return chunks

    def _determine_agent_type(self, file_path: str, text: str) -> str:
        """파일 경로와 내용을 기반으로 적절한 에이전트 타입을 결정합니다."""
        path_obj = Path(file_path)
        filename = path_obj.name.lower()
        text_lower = text.lower()

        # 1. 디렉토리 구조 기반 우선 분류
        # data/documents/source_code/ 형태의 경로에서 에이전트 타입 추출
        path_parts = path_obj.parts
        valid_agents = ['source_code', 'binary', 'log_conf']

        for part in path_parts:
            if part in valid_agents:
                print(f"📁 디렉토리 기반 분류: {file_path} → {part}")
                return part

        # 2. 파일명 기반 분류
        if any(keyword in filename for keyword in ['log', 'audit', 'config', 'conf', 'setting']):
            return 'log_conf'
        elif any(keyword in filename for keyword in ['binary', 'asm', 'disasm']):
            return 'binary'

        # 3. 내용 기반 분류
        source_code_keywords = ['function', 'class', 'import', 'include', 'def ', 'public class']
        binary_keywords = ['assembly', 'disassembly', 'binary analysis', 'executable']
        log_keywords = ['log', 'syslog', 'audit', 'connection', 'authentication', 'configuration', 'config', 'json', 'yaml', 'xml', 'settings']

        keyword_scores = {
            'source_code': sum(1 for kw in source_code_keywords if kw in text_lower),
            'binary': sum(1 for kw in binary_keywords if kw in text_lower),
            'log_conf': sum(1 for kw in log_keywords if kw in text_lower)
        }

        # 가장 높은 점수의 타입 선택
        best_type = max(keyword_scores.items(), key=lambda x: x[1])

        if best_type[1] > 0:
            print(f"📝 내용 기반 분류: {filename} → {best_type[0]} (점수: {best_type[1]})")
            return best_type[0]
        else:
            # 기본값: source_code
            print(f"🔧 기본값 분류: {filename} → source_code")
            return 'source_code'

    def get_document_summary(self, processed_doc: Dict[str, Any]) -> str:
        """문서 처리 결과 요약을 생성합니다."""
        info = processed_doc["file_info"]
        summary = f"""
📄 문서 처리 요약
- 파일명: {info['filename']}
- 크기: {info['file_size']:,} bytes
- 형식: {info['file_extension']}
- 에이전트 타입: {processed_doc['agent_type']}
- 총 청크 수: {processed_doc['total_chunks']}
- 원본 텍스트 길이: {processed_doc['original_text_length']:,} 문자
- 처리된 텍스트 길이: {processed_doc['processed_text_length']:,} 문자
"""
        return summary.strip()

# 의존성 주입을 위한 함수
def get_document_processor():
    return DocumentProcessor()